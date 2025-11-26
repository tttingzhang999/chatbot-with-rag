"""
Document processing service for RAG system.

Handles document upload, text extraction, chunking, embedding generation,
and BM25 index creation for hybrid search.
"""

import logging
import os
import re
from pathlib import Path
from typing import Any
from uuid import UUID

from sqlalchemy import delete, func, select, update
from sqlalchemy.orm import Session

from src.core.config import settings
from src.models.document import Document, DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Main document processing service.

    Orchestrates the complete pipeline: extract → chunk → embed → store
    """

    def __init__(self, db: Session):
        """Initialize processor with database session."""
        self.db = db
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

    def process_document_sync(
        self,
        document_id: UUID,
        file_path: str,
        file_type: str,
    ) -> dict[str, Any]:
        """
        Process document synchronously (simulates Lambda trigger).

        Args:
            document_id: UUID of the document record
            file_path: Path to the uploaded file
            file_type: File type (pdf, docx, txt)

        Returns:
            dict with processing results
        """
        try:
            logger.info(f"Starting document processing: {document_id}")

            # Update document status to processing
            self._update_document_status(document_id, "processing")

            # Step 1: Extract text
            text = self.extract_text(file_path, file_type)
            if not text or not text.strip():
                raise ValueError("No text content extracted from document")

            # Step 2: Chunk text
            chunks = self.chunk_text(text)
            logger.info(f"Created {len(chunks)} chunks from document {document_id}")

            # Step 3: Generate embeddings for all chunks
            embeddings = self.generate_embeddings(chunks)

            # Step 4: Create BM25 indexes
            bm25_vectors = self.create_bm25_indexes(chunks)

            # Step 5: Save chunks to database
            self.save_chunks_to_db(
                document_id=document_id,
                chunks=chunks,
                embeddings=embeddings,
                bm25_vectors=bm25_vectors,
            )

            # Update document status to completed
            self._update_document_status(document_id, "completed")

            logger.info(f"Document processing completed: {document_id}")
            return {
                "status": "success",
                "document_id": str(document_id),
                "chunks_created": len(chunks),
            }

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            self._update_document_status(
                document_id,
                "failed",
                error_message=str(e),
            )
            return {
                "status": "failed",
                "document_id": str(document_id),
                "error": str(e),
            }

    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from document.

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, txt)

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is unsupported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower()

        try:
            if file_type == "txt":
                # Read text file with UTF-8 encoding
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                logger.info(f"Extracted {len(text)} characters from TXT file")
                return text

            elif file_type == "pdf":
                # Extract text from PDF using pypdf
                # Filter out pages with garbled encoding (e.g., English pages with broken fonts)
                from pypdf import PdfReader

                reader = PdfReader(file_path)
                text_parts = []
                pages_processed = 0
                pages_skipped = 0

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if not page_text.strip():
                            continue

                        # Detect if page has readable Chinese content
                        # Skip pages with garbled encoding (low Chinese character ratio)
                        chinese_chars = sum(1 for c in page_text if "\u4e00" <= c <= "\u9fff")
                        total_chars = len(page_text.strip())

                        if total_chars > 0:
                            chinese_ratio = chinese_chars / total_chars

                            # Only include pages with >30% Chinese characters
                            # This filters out garbled English pages with broken font encoding
                            if chinese_ratio > 0.3:
                                text_parts.append(page_text)
                                pages_processed += 1
                            else:
                                pages_skipped += 1
                                logger.debug(
                                    f"Skipped page {page_num + 1} "
                                    f"(Chinese ratio: {chinese_ratio:.1%}, likely garbled encoding)"
                                )

                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue

                text = "\n\n".join(text_parts)
                logger.info(
                    f"Extracted {len(text)} characters from {pages_processed}/{len(reader.pages)} PDF pages "
                    f"({pages_skipped} pages skipped due to encoding issues)"
                )
                return text

            elif file_type in ["docx", "doc"]:
                # Extract text from DOCX using python-docx
                from docx import Document

                doc = Document(file_path)
                text_parts = []

                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)

                text = "\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from DOCX file")
                return text

            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except ImportError as e:
            logger.error(f"Missing dependency for {file_type} extraction: {e}")
            raise ValueError(
                f"Cannot extract {file_type} files. Missing required library. "
                f"Please install dependencies."
            )

    def chunk_text(
        self,
        text: str,
        chunk_size: int | None = None,
        overlap: int | None = None,
    ) -> list[str]:
        """
        Smart semantic chunking for structured and unstructured documents.

        Strategy:
        1. Detect if document is structured (has article/section markers)
        2. If structured: chunk by articles/sections
        3. If not: chunk by paragraphs
        4. Ensure chunks are within 300-1000 characters
        5. Add overlap for context preservation

        Args:
            text: Text to chunk
            chunk_size: Max chunk size in characters (default: 1000)
            overlap: Overlap size in characters (default: 100)

        Returns:
            List of semantically meaningful text chunks
        """
        # Set defaults for semantic chunking
        max_chunk_size = chunk_size or 1000
        min_chunk_size = 300
        overlap = overlap or 100

        if not text or not text.strip():
            return []

        # Normalize whitespace while preserving paragraph breaks
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()

        # Detect if document is structured
        is_structured = self._is_structured_document(text)

        if is_structured:
            logger.info("Detected structured document, chunking by articles/sections")
            semantic_units = self._split_by_articles(text)
        else:
            logger.info("Detected unstructured document, chunking by paragraphs")
            semantic_units = self._split_by_paragraphs(text)

        # Process semantic units into final chunks
        chunks = self._process_semantic_units(
            semantic_units,
            min_size=min_chunk_size,
            max_size=max_chunk_size,
            overlap=overlap,
        )

        logger.info(
            f"Created {len(chunks)} chunks from {len(semantic_units)} semantic units"
        )
        return chunks

    def _is_structured_document(self, text: str) -> bool:
        """
        Detect if document has structured markers (articles, sections).

        Args:
            text: Document text

        Returns:
            True if structured markers found
        """
        # Patterns for structured document markers
        patterns = [
            r"第[一二三四五六七八九十百]+條",  # Chinese: 第一條, 第二條
            r"Article\s+\d+",  # English: Article 1, Article 2
            r"Section\s+\d+",  # English: Section 1, Section 2
            r"第\d+條",  # Chinese with numbers: 第1條
        ]

        # Count marker occurrences
        marker_count = 0
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            marker_count += len(matches)

        # Consider structured if 3+ markers found
        return marker_count >= 3

    def _split_by_articles(self, text: str) -> list[str]:
        """
        Split text by article/section markers.

        Args:
            text: Document text

        Returns:
            List of article texts
        """
        # Combined pattern for all article markers
        article_pattern = r"(第[一二三四五六七八九十百\d]+條|Article\s+\d+|Section\s+\d+)"

        # Find all article positions
        articles = []
        matches = list(re.finditer(article_pattern, text, re.IGNORECASE))

        if not matches:
            # No markers found, return as single unit
            return [text]

        # Extract text between markers
        for i, match in enumerate(matches):
            start = match.start()
            # End is the start of next article or end of text
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

            article_text = text[start:end].strip()
            if article_text:
                articles.append(article_text)

        return articles

    def _split_by_paragraphs(self, text: str) -> list[str]:
        """
        Split text by paragraphs (double newline).

        Args:
            text: Document text

        Returns:
            List of paragraph texts
        """
        # Split by double newline (paragraph separator)
        paragraphs = re.split(r"\n\s*\n", text)

        # Clean and filter empty paragraphs
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        return paragraphs if paragraphs else [text]

    def _process_semantic_units(
        self,
        units: list[str],
        min_size: int,
        max_size: int,
        overlap: int,
    ) -> list[str]:
        """
        Process semantic units into final chunks with size constraints.

        Args:
            units: List of semantic units (articles/paragraphs)
            min_size: Minimum chunk size in characters
            max_size: Maximum chunk size in characters
            overlap: Overlap size for context

        Returns:
            List of processed chunks
        """
        chunks = []
        current_chunk = ""

        for unit in units:
            unit_size = len(unit)

            # Case 1: Unit fits within max_size
            if unit_size <= max_size:
                # If adding this unit exceeds max_size, save current chunk
                if current_chunk and len(current_chunk) + unit_size > max_size:
                    chunks.append(current_chunk.strip())
                    # Start new chunk with overlap from previous
                    current_chunk = self._get_overlap_text(current_chunk, overlap)

                # Add unit to current chunk
                current_chunk += "\n\n" + unit if current_chunk else unit

                # If current chunk meets min_size, consider saving it
                if len(current_chunk) >= min_size:
                    # Save if it's close to max_size or last unit
                    if len(current_chunk) >= max_size * 0.7 or unit == units[-1]:
                        chunks.append(current_chunk.strip())
                        current_chunk = ""

            # Case 2: Unit exceeds max_size, need to split further
            else:
                # Save any accumulated chunk first
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""

                # Split large unit by sentences
                sub_chunks = self._split_by_sentences(unit, max_size, overlap)
                chunks.extend(sub_chunks)

        # Add remaining chunk
        if current_chunk and len(current_chunk.strip()) > 0:
            chunks.append(current_chunk.strip())

        return [c for c in chunks if c]

    def _split_by_sentences(
        self, text: str, max_size: int, overlap: int
    ) -> list[str]:
        """
        Split text by sentences when it exceeds max_size.

        Args:
            text: Text to split
            max_size: Maximum chunk size
            overlap: Overlap size

        Returns:
            List of sentence-based chunks
        """
        # Split by sentence boundaries
        sentence_pattern = r"[。！？\.!\?]+\s*"
        sentences = re.split(sentence_pattern, text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current = ""

        for sentence in sentences:
            # If adding sentence exceeds max_size
            if current and len(current) + len(sentence) > max_size:
                chunks.append(current.strip())
                # Start new chunk with overlap
                current = self._get_overlap_text(current, overlap) + " " + sentence
            else:
                current += " " + sentence if current else sentence

        # Add remaining
        if current:
            chunks.append(current.strip())

        return chunks

    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """
        Get last N characters from text for overlap.

        Args:
            text: Source text
            overlap_size: Number of characters for overlap

        Returns:
            Overlap text
        """
        if len(text) <= overlap_size:
            return text

        # Get last overlap_size characters, try to break at word boundary
        overlap_text = text[-overlap_size:]

        # Find first space to avoid breaking words
        space_pos = overlap_text.find(" ")
        if space_pos > 0:
            overlap_text = overlap_text[space_pos:].strip()

        return overlap_text

    def generate_embeddings(self, chunks: list[str]) -> list[list[float]]:
        """
        Generate embeddings for text chunks using Cohere Embed v4 via Bedrock.

        Args:
            chunks: List of text chunks

        Returns:
            List of embedding vectors (1024 dimensions each)

        Raises:
            Exception: If embedding generation fails
        """
        from src.core.bedrock_client import get_bedrock_client

        try:
            bedrock_client = get_bedrock_client()
            embeddings = bedrock_client.generate_embeddings(
                texts=chunks, input_type="search_document"
            )

            logger.info(f"Generated {len(embeddings)} embeddings using Bedrock")
            return embeddings

        except Exception as e:
            logger.error(f"Failed to generate embeddings via Bedrock: {e}")
            raise

    def create_bm25_indexes(self, chunks: list[str]) -> list[str]:
        """
        Create BM25 full-text search indexes for chunks.

        Args:
            chunks: List of text chunks

        Returns:
            List of tsvector strings for PostgreSQL

        Note:
            PostgreSQL's to_tsvector() will be used in the database insert.
            This method prepares the text for indexing.
        """
        # Clean and prepare text for BM25 indexing
        # PostgreSQL's to_tsvector will handle the actual TSVECTOR creation
        processed_chunks = []

        for chunk in chunks:
            # Basic text cleaning for BM25
            # Remove special characters, normalize whitespace
            cleaned = re.sub(r"[^\w\s]", " ", chunk)
            cleaned = re.sub(r"\s+", " ", cleaned).strip().lower()
            processed_chunks.append(cleaned)

        return processed_chunks

    def save_chunks_to_db(
        self,
        document_id: UUID,
        chunks: list[str],
        embeddings: list[list[float]],
        bm25_vectors: list[str],
    ) -> None:
        """
        Save document chunks with embeddings and BM25 data to database.

        Args:
            document_id: UUID of the parent document
            chunks: List of text chunks
            embeddings: List of embedding vectors
            bm25_vectors: List of processed text for BM25
        """
        if not (len(chunks) == len(embeddings) == len(bm25_vectors)):
            raise ValueError("Chunks, embeddings, and BM25 vectors must have same length")

        try:
            # Create DocumentChunk objects
            chunk_objects = []
            for idx, (chunk_text, embedding, bm25_text) in enumerate(
                zip(chunks, embeddings, bm25_vectors, strict=True)
            ):
                chunk_obj = DocumentChunk(
                    document_id=document_id,
                    chunk_index=idx,
                    content=chunk_text,
                    embedding=embedding,
                    # Note: content_tsvector will be set by database trigger
                    # or we can use raw SQL: func.to_tsvector('english', chunk_text)
                    chunk_metadata={
                        "char_count": len(chunk_text),
                        "word_count": len(chunk_text.split()),
                    },
                )
                chunk_objects.append(chunk_obj)

            # Bulk insert chunks
            self.db.add_all(chunk_objects)
            self.db.commit()

            logger.info(f"Saved {len(chunk_objects)} chunks for document {document_id}")

        except Exception as e:
            self.db.rollback()
            logger.error(f"Error saving chunks to database: {e}")
            raise

    def _update_document_status(
        self,
        document_id: UUID,
        status: str,
        error_message: str | None = None,
    ) -> None:
        """
        Update document processing status.

        Args:
            document_id: Document UUID
            status: Status value (processing, completed, failed)
            error_message: Error message if status is failed
        """
        try:
            # Update document status
            update_values = {"status": status}
            if error_message is not None:
                update_values["error_message"] = error_message

            stmt = (
                update(Document)
                .where(Document.id == document_id)
                .values(**update_values)
            )
            self.db.execute(stmt)
            self.db.commit()

            logger.info(f"Updated document {document_id} status to: {status}")

        except Exception as e:
            self.db.rollback()
            logger.error(f"Error updating document status: {e}")
            # Don't raise here to avoid masking original errors


def get_user_documents(db: Session, user_id: UUID) -> list[dict[str, Any]]:
    """
    Get all documents uploaded by a user.

    Args:
        db: Database session
        user_id: User UUID

    Returns:
        List of document metadata dictionaries
    """
    stmt = (
        select(
            Document.id,
            Document.file_name,
            Document.file_type,
            Document.file_size,
            Document.upload_date,
            Document.status,
            Document.error_message,
            func.count(DocumentChunk.id).label("chunk_count"),
        )
        .where(Document.user_id == user_id)
        .outerjoin(DocumentChunk, Document.id == DocumentChunk.document_id)
        .group_by(
            Document.id,
            Document.file_name,
            Document.file_type,
            Document.file_size,
            Document.upload_date,
            Document.status,
            Document.error_message,
        )
        .order_by(Document.upload_date.desc())
    )

    result = db.execute(stmt)
    rows = result.all()

    documents = []
    for row in rows:
        documents.append(
            {
                "id": str(row.id),
                "file_name": row.file_name,
                "file_type": row.file_type,
                "file_size": row.file_size,
                "upload_date": row.upload_date.isoformat(),
                "status": row.status,
                "error_message": row.error_message,
                "chunk_count": row.chunk_count,
            }
        )

    return documents


def delete_document(db: Session, document_id: UUID, user_id: UUID) -> bool:
    """
    Delete a document and all its chunks.

    Args:
        db: Database session
        document_id: Document UUID
        user_id: User UUID (for authorization)

    Returns:
        True if deleted successfully, False if not found
    """
    try:
        # Check if document exists and belongs to user
        stmt = select(Document).where(
            Document.id == document_id,
            Document.user_id == user_id,
        )
        result = db.execute(stmt)
        document = result.scalar_one_or_none()

        if not document:
            return False

        # Delete associated file if it exists
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
                logger.info(f"Deleted file: {document.file_path}")
            except Exception as e:
                logger.warning(f"Failed to delete file {document.file_path}: {e}")

        # Delete document (cascades to chunks)
        stmt = delete(Document).where(Document.id == document_id)
        db.execute(stmt)
        db.commit()

        logger.info(f"Deleted document {document_id}")
        return True

    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting document {document_id}: {e}")
        raise
